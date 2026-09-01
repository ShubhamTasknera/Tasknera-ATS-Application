import io
import re
from docx import Document
from typing import Dict, Any
from app.services.text_cleaner import clean_extracted_text
from app.services.document_analyzer import analyze_document_quality

def parse_docx_bytes(docx_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Parses Word (.docx) documents using python-docx with fallback text recovery for malformed files.
    """
    extracted_text = ""
    extraction_method = "python-docx"

    try:
        doc = Document(io.BytesIO(docx_bytes))
        chunks = []

        # Extract text from paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                chunks.append(p.text.strip())

        # Extract text from tables without missing cell data
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_texts:
                    chunks.append(" | ".join(row_texts))

        extracted_text = "\n".join(chunks)
    except Exception as e:
        extraction_method = "docx-stream-recovery"
        # Binary text recovery fallback for corrupted or legacy .doc containers
        try:
            raw_str = docx_bytes.decode('utf-8', errors='ignore')
            clean_printable = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', raw_str)
            # Find runs of alphanumeric characters
            words = re.findall(r'[A-Za-z0-9@+.,\-/\s]{4,}', clean_printable)
            extracted_text = "\n".join([w.strip() for w in words if len(w.strip()) > 3])
        except Exception:
            extracted_text = ""

    cleaned_text = clean_extracted_text(extracted_text)
    metrics = analyze_document_quality(cleaned_text, 1, filename)

    return {
        "text": cleaned_text,
        "pageCount": 1,
        "extractionMethod": extraction_method,
        "ocrUsed": False,
        "textQuality": metrics["textQuality"],
        "characterCount": metrics["characterCount"],
        "wordCount": metrics["wordCount"],
        # Structured Fields
        "candidateName": metrics["candidateName"],
        "email": metrics["email"],
        "phone": metrics["phone"],
        "skills": metrics["skills"],
        "yearsOfExperience": metrics["yearsOfExperience"],
        "education": metrics["education"],
        "pastCompanies": metrics["pastCompanies"],
        "summary": metrics["summary"],
        "rawTextSummary": metrics["rawTextSummary"],
    }

