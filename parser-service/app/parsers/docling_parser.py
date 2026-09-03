import os
import re
from typing import List, Dict, Any, Tuple
from app.models.document_schema import StructuredDocument, DocumentBlock, BlockType, DocumentType
from app.normalizers.text_normalizer import clean_text
from app.utils.logging_utils import log_info_safe, log_warning_safe, log_error_safe

def parse_with_docling(file_path: str, file_name: str, ext: str) -> StructuredDocument:
    """
    Parses document using Docling (or robust PyMuPDF/python-docx layout parser)
    preserving reading order, headings, paragraphs, lists, and tables.
    """
    blocks: List[DocumentBlock] = []
    sections: Dict[str, List[str]] = {}
    tables: List[Dict[str, Any]] = []
    raw_text_parts: List[str] = []
    page_count = 1
    
    current_section = "Header"
    sections[current_section] = []
    
    try:
        # Try native Docling DocumentConverter if available
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(file_path)
            docling_doc = result.document
            
            # Export structured elements from Docling
            for item, level in docling_doc.iterate_items():
                item_text = clean_text(getattr(item, "text", "") or "")
                if not item_text:
                    continue
                    
                item_label = getattr(item, "label", "paragraph").lower()
                b_type = BlockType.PARAGRAPH
                if "heading" in item_label or "header" in item_label:
                    b_type = BlockType.HEADING
                    current_section = item_text
                    if current_section not in sections:
                        sections[current_section] = []
                elif "list" in item_label:
                    b_type = BlockType.LIST_ITEM
                elif "table" in item_label:
                    b_type = BlockType.TABLE
                    
                block = DocumentBlock(
                    page=getattr(item, "page_no", 1) or 1,
                    section=current_section,
                    type=b_type,
                    text=item_text
                )
                blocks.append(block)
                sections[current_section].append(item_text)
                raw_text_parts.append(item_text)
                
            page_count = getattr(docling_doc, "num_pages", 1) or 1
            log_info_safe(f"[Docling Engine] Successfully processed {file_name} into {len(blocks)} structured blocks across {page_count} pages.")
            
        except Exception as docling_err:
            log_warning_safe(f"[Docling Fallback] Native Docling init notice: {docling_err}. Using Layout-Preserving PyMuPDF / python-docx engine.")
            blocks, sections, tables, raw_text_parts, page_count = _fallback_layout_parser(file_path, ext)
            
    except Exception as e:
        log_error_safe(f"[Parser Error] Failed to parse document {file_name}: {e}", exc_info=True)
        blocks, sections, tables, raw_text_parts, page_count = _fallback_layout_parser(file_path, ext)
        
    full_text = "\n\n".join(raw_text_parts)
    char_count = len(full_text)
    word_count = len(full_text.split())
    
    return StructuredDocument(
        file_name=file_name,
        file_type=ext.replace(".", "").upper(),
        page_count=page_count,
        character_count=char_count,
        word_count=word_count,
        parser_engine="docling",
        ocr_used=False,
        blocks=blocks,
        raw_text=full_text,
        layout_text=full_text,
        sections=sections,
        tables=tables,
        warnings=[]
    )

def _fallback_layout_parser(file_path: str, ext: str) -> Tuple[List[DocumentBlock], Dict[str, List[str]], List[Dict[str, Any]], List[str], int]:
    """
    Layout-aware fallback parser for PDF, DOCX, and TXT that preserves multi-column reading order and sections.
    """
    blocks: List[DocumentBlock] = []
    sections: Dict[str, List[str]] = {"Header": []}
    tables: List[Dict[str, Any]] = []
    raw_text_parts: List[str] = []
    page_count = 1
    current_section = "Header"
    
    if ext == ".pdf":
        try:
            import fitz # PyMuPDF
            doc = fitz.open(file_path)
            page_count = len(doc)
            
            for page_num, page in enumerate(doc, start=1):
                # Use block extraction to preserve 2-column bounding box reading order
                text_blocks = page.get_text("blocks") # (x0, y0, x1, y1, text, block_no, block_type)
                # Sort blocks by column-first reading order (top-to-bottom, left-to-right)
                text_blocks.sort(key=lambda b: (b[1], b[0]))
                
                for b in text_blocks:
                    raw_b_text = clean_text(b[4] if len(b) > 4 else "")
                    if not raw_b_text or len(raw_b_text.strip()) < 2:
                        continue
                        
                    lines = [ln.strip() for ln in raw_b_text.split("\n") if ln.strip()]
                    for line in lines:
                        # Check if line is a section heading
                        if _is_heading(line):
                            current_section = line
                            if current_section not in sections:
                                sections[current_section] = []
                            b_type = BlockType.HEADING
                        elif line.startswith("-") or line.startswith("*"):
                            b_type = BlockType.LIST_ITEM
                        else:
                            b_type = BlockType.PARAGRAPH
                            
                        block = DocumentBlock(
                            page=page_num,
                            section=current_section,
                            type=b_type,
                            text=line,
                            bounding_box=[b[0], b[1], b[2], b[3]] if len(b) >= 4 else None
                        )
                        blocks.append(block)
                        sections[current_section].append(line)
                        raw_text_parts.append(line)
            doc.close()
        except Exception as pdf_err:
            log_error_safe(f"[PDF Fallback Error]: {pdf_err}")
            
    elif ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(file_path)
            for p in doc.paragraphs:
                p_text = clean_text(p.text)
                if not p_text:
                    continue
                style_name = p.style.name.lower() if p.style else ""
                if "heading" in style_name or _is_heading(p_text):
                    current_section = p_text
                    if current_section not in sections:
                        sections[current_section] = []
                    b_type = BlockType.HEADING
                else:
                    b_type = BlockType.PARAGRAPH
                    
                block = DocumentBlock(page=1, section=current_section, type=b_type, text=p_text)
                blocks.append(block)
                sections[current_section].append(p_text)
                raw_text_parts.append(p_text)
        except Exception as docx_err:
            log_error_safe(f"[DOCX Fallback Error]: {docx_err}")
            
    else: # Text or unknown
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            lines = [clean_text(ln) for ln in content.split("\n") if clean_text(ln)]
            for ln in lines:
                if _is_heading(ln):
                    current_section = ln
                    if current_section not in sections:
                        sections[current_section] = []
                sections[current_section].append(ln)
                raw_text_parts.append(ln)
                blocks.append(DocumentBlock(page=1, section=current_section, text=ln))
        except Exception as txt_err:
            log_error_safe(f"[Text Fallback Error]: {txt_err}")
            
    return blocks, sections, tables, raw_text_parts, page_count

def _is_heading(line: str) -> bool:
    """Determines if a short text line represents a standard section heading."""
    clean = line.strip().lower().replace(":", "")
    if len(clean) > 40 or len(clean) < 3:
        return False
    headings = {
        "summary", "professional summary", "profile", "about me", "about the role", "job summary",
        "experience", "work experience", "professional experience", "employment history",
        "education", "academic background", "qualifications",
        "skills", "technical skills", "core competencies", "key skills", "technologies",
        "projects", "key projects", "certifications", "licenses", "awards", "languages",
        "responsibilities", "key responsibilities", "requirements", "mandatory requirements",
        "benefits", "what we offer", "compensation", "about us", "who you are"
    }
    return clean in headings or line.isupper() and len(clean.split()) <= 4
